#!/usr/bin/env python3
"""Генерация docx из markdown-файлов требований и ТК (v2.1 / v1.2)."""
from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DESKTOP = Path(r"c:\Users\User\Desktop\Проект эндуро")

FILES = [
    (
        DOCS / "SYSTEM-REQUIREMENTS-v2.1.md",
        [
            DESKTOP / "Системные требования к проекту Enduro Park Manager v2.2.docx",
            DESKTOP / "Системные требования к проекту Enduro Park Manager.docx",
        ],
    ),
    (
        DOCS / "MANUAL-TEST-CASES-v1.2.md",
        [
            DESKTOP / "Тестовые кейсы ручные v1.3.docx",
            DESKTOP / "Тестовые кейсы ручные.docx",
        ],
    ),
]


def add_markdown_to_doc(doc: Document, text: str) -> None:
    for line in text.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("|") and "---" not in line:
            p = doc.add_paragraph(line.strip())
            p.style = "Normal"
            for run in p.runs:
                run.font.size = Pt(10)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(11)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    DESKTOP.mkdir(parents=True, exist_ok=True)

    for src, targets in FILES:
        if not src.exists():
            print(f"Skip missing: {src}")
            continue
        text = src.read_text(encoding="utf-8")
        doc = Document()
        add_markdown_to_doc(doc, text)
        for target in targets:
            doc.save(str(target))
            print(f"Written: {target}")


if __name__ == "__main__":
    main()
